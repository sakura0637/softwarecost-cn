import re, sys, traceback

PDF = "C:/Users/lenovo/Desktop/标准(2)/标准/山西省省直部门信息化建设项目支出预算方案编制规范和预算编制标准.pdf"
DOCX36964 = "C:/Users/lenovo/Desktop/预算编制/标准/GBT+36964-2018.docx"
DOCX28827 = "C:/Users/lenovo/Desktop/预算编制/标准/GBT+28827.7-2022(1).docx"

KEYWORDS = ["人月", "费率", "功能点", "UFP", "ILF", "EIF", "EI", "EO", "EQ", "调整因子",
            "生产率", "运维", "开发", "城市", "权重", "VAF", "复用", "规模", "核算",
            "元/人月", "万元", "基准", "成本", "折算"]

def banner(t): print("\n" + "="*70 + "\n" + t + "\n" + "="*70)

def probe_pdf(path, title):
    banner(title)
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            print("页数:", len(pdf.pages))
            full = []
            for i, page in enumerate(pdf.pages):
                full.append(page.extract_text() or "")
            txt = "\n".join(full)
            print("总字符数:", len(txt))
            for kw in KEYWORDS:
                c = txt.count(kw)
                if c: print(f"  {kw}: {c}")
            print("\n--- 含'费率'/'人月'/'功能点'的片段(前30条) ---")
            cnt = 0
            for line in txt.splitlines():
                if re.search(r"费率|人月|功能点|调整因子", line):
                    print(line.strip()[:130])
                    cnt += 1
                    if cnt >= 30: break
            # tables containing key terms
            print("\n--- 可能含参数的表格(每表前6行) ---")
            for i, page in enumerate(pdf.pages):
                for ti, tb in enumerate(page.extract_tables() or []):
                    flat = "\n".join(" | ".join(str(c or "") for c in r) for r in tb)
                    if re.search(r"费率|人月|功能点|ILF|EIF|调整|权重|元", flat):
                        print(f"\n[页{i+1} 表{ti}]")
                        for r in tb[:6]:
                            print(" | ".join(str(c or "") for c in r)[:130])
    except Exception as e:
        traceback.print_exc()

def probe_docx(path, title):
    banner(title)
    try:
        import docx
        d = docx.Document(path)
        paras = [p.text for p in d.paragraphs]
        txt = "\n".join(paras)
        print("段落数:", len(paras), " 总字符:", len(txt))
        for kw in KEYWORDS:
            c = txt.count(kw)
            if c: print(f"  {kw}: {c}")
        print("\n--- 含关键字的段落(前30条) ---")
        cnt = 0
        for p in paras:
            if re.search(r"功能点|权重|VAF|调整因子|复杂度|ILF|EIF|EI|EO|EQ|人月|费率|生产率|运维|开发", p):
                print(p.strip()[:150]); cnt += 1
                if cnt >= 30: break
        print("\n表格数:", len(d.tables))
        for ti, tb in enumerate(d.tables):
            rows = [" | ".join(c.text.strip() for c in r.cells) for r in tb.rows]
            flat = "\n".join(rows)
            if re.search(r"ILF|EIF|EI|EO|EQ|权重|复杂度|调整因子|人月|费率|生产率|元|开发|运维", flat):
                print(f"\n[表{ti}] 可能含参数:")
                for r in rows[:8]:
                    print("   ", r[:150])
    except Exception as e:
        traceback.print_exc()

probe_pdf(PDF, "① 山西省省直部门信息化建设项目支出预算方案编制规范 (PDF)")
probe_docx(DOCX36964, "② GB/T 36964-2018 (DOCX)")
probe_docx(DOCX28827, "③ GB/T 28827.7-2022 运维成本度量规范 (DOCX)")
