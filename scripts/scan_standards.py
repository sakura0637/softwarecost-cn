#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""扫描标准 PDF/DOCX，抽取与造价页面相关的关键参数，输出紧凑摘要。"""
import sys, os, re, glob
from pathlib import Path

PDF_DIRS = [
    r"C:/Users/lenovo/Desktop/标准(2)/标准",
    r"C:/Users/lenovo/Desktop/预算编制/标准",
]

# 关键词：命中即说明该文件含对应数据
KEYWORDS = {
    "城市级费率": r"城市|青岛|长沙|重庆|成都|广州|北京|上海|深圳|杭州|南京|武汉|西安|天津|苏州|郑州|济南|合肥|沈阳|大连|厦门|宁波|石家庄|太原|福州|南昌|昆明|贵阳|南宁|兰州|海口|呼和浩特|银川|西宁|乌鲁木齐",
    "元/人月": r"元/人月|元/?人月|人月费率|人月单价",
    "元/FP": r"元/FP|元每功能点|功能点单价",
    "生产率P": r"P10|P25|P50|P75|P90|生产率",
    "UFP权重": r"ILF|EIF|EI\b|EO\b|EQ\b|功能点计数|未调整功能点|复杂度",
    "调整因子": r"调整因子|规模变更|应用类型|开发平台|开发团队|非功能|完整性级别|软件规模",
    "人月折算": r"人月折算系数|HM\b|21\.75|折算系数",
    "人力费率": r"人力成本费率|平均人力|人月费用单价|人工单价|费用单价",
    "运维": r"运维|运行维护|维护费用|OPEX|运维费率",
    "年份": r"201[89]|202[0-5]|近[一二三]年",
}

def extract_pdf_text(path, max_chars=60000):
    import pdfplumber
    out = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            if sum(len(x) for x in out) > max_chars:
                break
            t = page.extract_text() or ""
            out.append(t)
    return "\n".join(out)

def extract_docx_text(path, max_chars=60000):
    import docx
    d = docx.Document(path)
    out = []
    for p in d.paragraphs:
        out.append(p.text)
    for t in d.tables:
        for row in t.rows:
            out.append(" | ".join(c.text for c in row.cells))
    return "\n".join(out)

def scan():
    files = []
    for d in PDF_DIRS:
        files += glob.glob(os.path.join(d, "*.pdf"))
        files += glob.glob(os.path.join(d, "*.doc"))
        files += glob.glob(os.path.join(d, "*.docx"))
    files = sorted(set(files))
    print(f"共 {len(files)} 个文件\n{'='*70}")
    for f in files:
        name = os.path.basename(f)
        try:
            if f.lower().endswith('.pdf'):
                text = extract_pdf_text(f)
            else:
                text = extract_docx_text(f)
        except Exception as e:
            print(f"[读取失败] {name}: {e}")
            continue
        if not text.strip():
            print(f"[空] {name}")
            continue
        hits = {}
        for k, pat in KEYWORDS.items():
            n = len(re.findall(pat, text))
            if n:
                hits[k] = n
        # 抽取几个样例行（含 元/人月 或 费率 的）
        sample = []
        for line in text.splitlines():
            if re.search(r"元/人月|人月费率|人月单价|元/FP|生产率|P50|调整因子|人月折算|人力成本", line) and len(line) < 80:
                sample.append(line.strip())
            if len(sample) >= 6:
                break
        print(f"\n📄 {name}  ({len(text)} 字)")
        print("   命中: " + ", ".join(f"{k}×{v}" for k, v in hits.items()) or "  (无关键词)")
        if sample:
            for s in sample[:5]:
                print(f"   · {s}")

if __name__ == "__main__":
    scan()
